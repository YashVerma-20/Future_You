"""DOCX text extraction utility."""
import io
from typing import Optional
from docx import Document
import structlog

logger = structlog.get_logger()


def extract_text_from_docx(file_content: bytes) -> Optional[str]:
    """
    Extract text from DOCX file content.
    
    Args:
        file_content: DOCX file content as bytes
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        full_text = "\n".join(text)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        return full_text.strip() if full_text else None
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return None


def extract_docx_metadata(file_content: bytes) -> dict:
    """
    Extract metadata from DOCX file.
    
    Args:
        file_content: DOCX file content as bytes
        
    Returns:
        Dictionary containing DOCX metadata
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        core_props = doc.core_properties
        
        metadata = {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'created': core_props.created.isoformat() if core_props.created else None,
            'modified': core_props.modified.isoformat() if core_props.modified else None,
            'paragraph_count': len(doc.paragraphs),
        }
        
        return metadata
        
    except Exception as e:
        logger.error(f"DOCX metadata extraction failed: {e}")
        return {}
